import os
import re
import threading
from markitdown import MarkItDown
from llama_cloud_services import LlamaParse
from llama_cloud_services.parse.utils import ParsingMode
from app.core.config import settings
from docx import Document

class FileProcessor:
    _llama_instance = None
    _llama_lock = threading.Lock()

    def __init__(self):
        # MarkItDown có thể tạo mỗi instance riêng cho mỗi lần gọi
        self.md = MarkItDown(enable_plugins=False)


    def _clean_text(self, text: str) -> str:
        """Xóa các cột Unnamed và NaN"""
        text = re.sub(r"Unnamed:\s*\d+", "", text)
        text = re.sub(r"\bNaN\b", "", text)
        return text

    def _split_sheets(self, markdown_text: str) -> list[str]:
        """Tách văn bản markdown thành các sheet theo header ##"""
        sheets = re.split(r"(?=^## )", markdown_text, flags=re.MULTILINE)
        return [part.strip() for part in sheets if part.strip().startswith("##")]

    def read_excel(self, file_path: str) -> list[str]:
        """Đọc file Excel, chuyển thành Markdown, dọn sạch, rồi tách sheet"""
        markdown_text = self.md.convert(file_path).text_content
        markdown_text_clean = self._clean_text(markdown_text)
        return self._split_sheets(markdown_text_clean)

    def _docx_has_image_or_table(self, file_path: str) -> bool:
        doc = Document(file_path)
        # Check for tables
        if doc.tables:
            return True
        # Check for images (inline shapes)
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                # if 'w:drawing' in run._element.xml:
                if run._element.xpath('.//pic:pic'):
                    return True
        return False

    def read_markdownable(self, file_path: str) -> str:
        """Đọc file text-like (md, txt)"""
        file_ext = file_path.split(".")[-1].lower()
        if file_ext == "docx":
            if self._docx_has_image_or_table(file_path):
                print("DOCX contains image or table, using LlamaParse...")
                parser = self._get_llama_parser()
                parse_result = parser.parse(file_path)
                markdown_documents = parse_result.get_markdown_documents()
                text = "\n\n".join(doc.text_resource.text for doc in markdown_documents)
                return text
            else:
                print("DOCX contains only text, using MarkItDown...")
                return self.md.convert(file_path).text_content
        else:
            print("Processing non-docx file with MarkItDown...")
            return self.md.convert(file_path).text_content

    def read_with_llama(self, file_path: str) -> str:
        """Đọc file PDF, ảnh"""
        print("Processing pdf with LlamaParse...")
        parser = self._get_llama_parser()
        parse_result = parser.parse(file_path)
        markdown_documents = parse_result.get_markdown_documents()
        text = "\n\n".join(doc.text_resource.text for doc in markdown_documents)
        return text

        
    def read_file(self, file_path: str) -> str | list[str]:
        """Router: chọn cách đọc file tùy loại"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        file_ext = file_path.split(".")[-1].lower()
        if file_ext in ["xlsx", "xls"]:
            return self.read_excel(file_path)
        elif file_ext in ["pdf", "jpg", "jpeg", "png", "gif", "webp", "bmp", "tiff", "tif"]:
            return self.read_with_llama(file_path)
        else:
            return self.read_markdownable(file_path)

file_processor = FileProcessor()

# Ví dụ sử dụng
if __name__ == "__main__":
    processor = FileProcessor()
    import time
    start_time = time.time()    
    sheets = processor.read_file(rf"c:\Users\four\Downloads\data-ai-marketing_2609_v3\show_daily(15).pdf")
    # print(sheets[0])
    with open('read_file.md2', "w", encoding="utf-8") as f:
        if isinstance(sheets, list):
            for sheet in sheets:
                f.write(sheet + "\n\n")
        else:
            f.write(sheets)
    print(time.time()- start_time)

    # with open(r"C:\Users\Admin\Downloads\test_2.md", "w", encoding="utf-8") as f:
    #     f.write(sheets[0])